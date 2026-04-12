"""Tests for DOCX and Markdown export."""

import io
from docx import Document
from docx.shared import Pt

from src.exporter import export_docx, export_markdown
from src.models import ExportRequest


def _sample_request() -> ExportRequest:
    return ExportRequest(
        title="Widget Manufacturing System",
        cross_references="",
        background="The field of widget manufacturing has grown.\n\nExisting methods are slow.",
        summary="The present invention provides a faster widget system.",
        detailed_description="In a preferred embodiment, the system comprises a hopper.\n\nThe hopper feeds into a **conveyor belt** that operates at *high speed*.",
        claims="1. A method of manufacturing widgets comprising a hopper and a conveyor.\n\n2. The method of claim 1, further comprising a sorter.",
        abstract="A widget manufacturing system with improved throughput.",
        figure_descriptions="FIG. 1 is a block diagram showing the widget system.",
        ids_table="| Ref | Patent Number | Title |\n|-----|-------------|-------|\n| 1 | US10123456 | Old Widget |",
    )


def _load_docx(data: bytes) -> Document:
    """Load a Document from raw bytes for inspection."""
    return Document(io.BytesIO(data))


class TestDocxExport:
    def test_produces_valid_zip_bytes(self):
        result = export_docx(_sample_request())
        assert isinstance(result, bytes)
        assert len(result) > 0
        # DOCX files are ZIP archives — PK magic header
        assert result[:2] == b"PK"

    def test_empty_sections_handled(self):
        req = ExportRequest(title="Minimal", claims="1. A method.")
        result = export_docx(req)
        assert isinstance(result, bytes)
        assert len(result) > 0
        assert result[:2] == b"PK"

    def test_times_new_roman_font(self):
        result = export_docx(_sample_request())
        doc = _load_docx(result)
        style = doc.styles["Normal"]
        assert style.font.name == "Times New Roman"
        assert style.font.size == Pt(12)

    def test_line_spacing_1_5(self):
        result = export_docx(_sample_request())
        doc = _load_docx(result)
        style = doc.styles["Normal"]
        assert style.paragraph_format.line_spacing == 1.5

    def test_page_margins(self):
        result = export_docx(_sample_request())
        doc = _load_docx(result)
        section = doc.sections[0]
        # Inches are stored as EMU — 914400 EMU per inch
        emu_per_inch = 914400
        assert section.left_margin == 1.0 * emu_per_inch
        assert section.right_margin == 0.75 * emu_per_inch
        assert section.top_margin == 0.75 * emu_per_inch
        assert section.bottom_margin == 0.75 * emu_per_inch

    def test_page_size_letter(self):
        result = export_docx(_sample_request())
        doc = _load_docx(result)
        section = doc.sections[0]
        emu_per_inch = 914400
        assert section.page_width == 8.5 * emu_per_inch
        assert section.page_height == 11 * emu_per_inch

    def test_paragraph_numbering_present(self):
        """Numbered sections should contain [NNNN] prefixed paragraphs."""
        result = export_docx(_sample_request())
        doc = _load_docx(result)
        texts = [p.text for p in doc.paragraphs]
        numbered = [t for t in texts if t.startswith("[0001]")]
        assert len(numbered) >= 1

    def test_paragraph_numbering_continuous(self):
        """Numbers should be continuous across Background, Summary, Figures, DD."""
        result = export_docx(_sample_request())
        doc = _load_docx(result)
        nums = []
        for p in doc.paragraphs:
            text = p.text
            if text.startswith("[") and "]" in text[:6]:
                num_str = text[1 : text.index("]")]
                if num_str.isdigit():
                    nums.append(int(num_str))
        # Should be 1, 2, 3, ... with no gaps
        assert nums == list(range(1, len(nums) + 1))

    def test_claims_not_numbered(self):
        """Claims section should NOT have [NNNN] numbering."""
        result = export_docx(_sample_request())
        doc = _load_docx(result)
        in_claims = False
        for p in doc.paragraphs:
            if "CLAIMS" in p.text.upper() and p.alignment is not None:
                in_claims = True
                continue
            if in_claims and p.text.upper().startswith("ABSTRACT"):
                break
            if in_claims and p.text.strip():
                assert not p.text.startswith("[0")

    def test_abstract_not_numbered(self):
        """Abstract should NOT have [NNNN] numbering."""
        result = export_docx(_sample_request())
        doc = _load_docx(result)
        in_abstract = False
        for p in doc.paragraphs:
            if "ABSTRACT" in p.text.upper() and p.alignment is not None:
                in_abstract = True
                continue
            if in_abstract and p.text.strip():
                assert not p.text.startswith("[0")

    def test_page_numbers_in_footer(self):
        """Footer should contain PAGE field code for page numbers."""
        result = export_docx(_sample_request())
        doc = _load_docx(result)
        section = doc.sections[0]
        footer_xml = section.footer._element.xml
        assert "PAGE" in footer_xml

    def test_footer_disclaimer_text(self):
        """Footer should contain the PatentForge disclaimer."""
        result = export_docx(_sample_request())
        doc = _load_docx(result)
        section = doc.sections[0]
        footer_text = "\n".join(p.text for p in section.footer.paragraphs)
        assert "PatentForge" in footer_text

    def test_watermark_in_header(self):
        """Header should contain draft warning banner text."""
        result = export_docx(_sample_request())
        doc = _load_docx(result)
        section = doc.sections[0]
        header_text = "\n".join(p.text for p in section.header.paragraphs)
        assert "NOT LEGAL ADVICE" in header_text
        assert "PATENT ATTORNEY" in header_text

    def test_markdown_bold_rendered(self):
        """**bold** markers should be converted to actual bold runs."""
        result = export_docx(_sample_request())
        doc = _load_docx(result)
        # Find paragraph containing "conveyor belt" (was **conveyor belt**)
        found_bold = False
        for p in doc.paragraphs:
            for run in p.runs:
                if "conveyor belt" in run.text and run.bold:
                    found_bold = True
                    break
        assert found_bold, "Expected **conveyor belt** to render as bold run"

    def test_markdown_italic_rendered(self):
        """*italic* markers should be converted to actual italic runs."""
        result = export_docx(_sample_request())
        doc = _load_docx(result)
        found_italic = False
        for p in doc.paragraphs:
            for run in p.runs:
                if "high speed" in run.text and run.italic:
                    found_italic = True
                    break
        assert found_italic, "Expected *high speed* to render as italic run"

    def test_strips_existing_paragraph_numbers(self):
        """If input text already has [NNNN] prefixes, they should be stripped."""
        req = ExportRequest(
            title="Test",
            background="[0001] Already numbered paragraph.\n\n[0002] Second numbered.",
        )
        result = export_docx(req)
        doc = _load_docx(result)
        # Should have [0001] and [0002] from the exporter, not doubled
        texts = [p.text for p in doc.paragraphs]
        numbered = [t for t in texts if "[0001]" in t]
        assert len(numbered) == 1


class TestMarkdownExport:
    def test_produces_string(self):
        result = export_markdown(_sample_request())
        assert isinstance(result, str)
        assert "Widget Manufacturing System" in result
        assert "## Background" in result or "Background of the Invention" in result
        assert "## Claims" in result or "Claims" in result

    def test_includes_ids_table(self):
        result = export_markdown(_sample_request())
        assert "Information Disclosure Statement" in result
        assert "US10123456" in result

    def test_empty_sections_omitted(self):
        req = ExportRequest(title="Minimal", claims="1. A method.")
        result = export_markdown(req)
        assert "Claims" in result
        assert "Background" not in result
